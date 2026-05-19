from __future__ import annotations

import io
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from uuid import UUID, uuid7

import pytesseract
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile
from PIL import Image
from sqlmodel import Session, select

from app.agents.deps import AgentDeps
from app.agents.document import build_document_chat_agent
from app.core.security import decrypt_string
from app.models.document import Document
from app.models.document_chunk import DocumentChunk
from app.models.user import User
from app.services.ai_service import get_embedding

MAX_DOCUMENT_BYTES = int(os.getenv("DOCUMENT_MAX_BYTES", str(25 * 1024 * 1024)))
CHUNK_SIZE = int(os.getenv("DOCUMENT_CHUNK_SIZE", "1200"))
CHUNK_OVERLAP = int(os.getenv("DOCUMENT_CHUNK_OVERLAP", "200"))
STORAGE_ROOT = os.getenv("DOCUMENT_STORAGE_PATH", "uploads")
OCR_LANG = os.getenv("OCR_LANG", "eng")


def _resolve_api_key(current_user: User) -> str | None:
    if current_user.encrypted_google_api_key:
        return decrypt_string(current_user.encrypted_google_api_key)

    env_key = os.getenv("GOOGLE_API_KEY", "")
    return env_key or None


def _ensure_storage_dir(user_id: UUID) -> Path:
    root = Path(STORAGE_ROOT)
    user_dir = root / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir


def _normalize_text(text: str) -> str:
    return " ".join(text.split()).strip()


def _chunk_text(text: str, max_chars: int, overlap: int) -> list[str]:
    normalized = _normalize_text(text)
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    length = len(normalized)
    while start < length:
        end = min(start + max_chars, length)
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= length:
            break
        start = max(0, end - overlap)

    return chunks


def _extract_text_from_image(data: bytes) -> str:
    image = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(image, lang=OCR_LANG)


def _extract_text_from_pdf(data: bytes) -> str:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    texts: list[str] = []

    for page in doc:
        page_text = page.get_text().strip()
        if page_text:
            texts.append(page_text)
            continue

        pix = page.get_pixmap(dpi=200)
        mode = "RGBA" if pix.alpha else "RGB"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        if pix.alpha:
            img = img.convert("RGB")
        ocr_text = pytesseract.image_to_string(img, lang=OCR_LANG)
        if ocr_text.strip():
            texts.append(ocr_text)

    return "\n\n".join(texts)


def _extract_text_from_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _extract_text_from_plain(data: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _detect_type(filename: str, content_type: str | None) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or content_type == "application/pdf":
        return "pdf"

    if content_type and content_type.startswith("image/"):
        return "image"

    if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif", ".gif"}:
        return "image"

    if ext in {".docx"}:
        return "docx"

    if ext in {".txt", ".md", ".csv", ".log"}:
        return "text"

    return "unknown"


def _extract_text(filename: str, content_type: str | None, data: bytes) -> str:
    doc_type = _detect_type(filename, content_type)

    if doc_type == "pdf":
        return _extract_text_from_pdf(data)

    if doc_type == "image":
        return _extract_text_from_image(data)

    if doc_type == "docx":
        return _extract_text_from_docx(data)

    if doc_type == "text":
        return _extract_text_from_plain(data)

    raise HTTPException(status_code=400, detail="Unsupported document type")


async def create_document_service(
    session: Session,
    current_user: User,
    upload: UploadFile,
    title: Optional[str] = None,
) -> Document:
    filename = upload.filename or "document"
    data = await upload.read()
    await upload.close()

    if len(data) > MAX_DOCUMENT_BYTES:
        raise HTTPException(status_code=400, detail="Document too large")

    storage_dir = _ensure_storage_dir(current_user.id)
    ext = Path(filename).suffix
    stored_name = f"{uuid7()}{ext}"
    storage_path = storage_dir / stored_name
    storage_path.write_bytes(data)

    extracted_text = _extract_text(filename, upload.content_type, data)
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="No text extracted from document")

    chunks = _chunk_text(extracted_text, CHUNK_SIZE, CHUNK_OVERLAP)
    if not chunks:
        raise HTTPException(status_code=400, detail="No text chunks generated")

    api_key = _resolve_api_key(current_user)
    if current_user.ai_provider == "google" and not api_key:
        raise HTTPException(status_code=400, detail="Google API Key not configured")

    document = Document(
        user_id=current_user.id,
        title=title or filename,
        filename=filename,
        content_type=upload.content_type or "application/octet-stream",
        storage_path=str(storage_path),
        text=extracted_text,
        metadata_info={
            "chunk_count": str(len(chunks)),
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
        },
    )

    try:
        session.add(document)
        session.flush()

        for index, chunk in enumerate(chunks):
            embedding = await get_embedding(
                chunk,
                is_search=False,
                provider=current_user.ai_provider,
                api_key=api_key,
            )
            session.add(
                DocumentChunk(
                    document_id=document.id,
                    user_id=current_user.id,
                    chunk_index=index,
                    content=chunk,
                    embedding=embedding,
                )
            )

        session.commit()
        session.refresh(document)
        return document
    except Exception as exc:
        session.rollback()
        raise HTTPException(
            status_code=500, detail=f"Failed to store document: {str(exc)}"
        )


def list_documents_service(
    session: Session,
    current_user: User,
    limit: int = 50,
    offset: int = 0,
) -> list[Document]:
    statement = (
        select(Document)
        .where(Document.user_id == current_user.id)
        .order_by(Document.created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    return list(session.exec(statement).all())


def get_document_service(
    session: Session,
    current_user: User,
    document_id: UUID,
) -> Document:
    document = session.get(Document, document_id)
    if not document or document.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Document not found")
    return document


async def search_document_chunks(
    session: Session,
    current_user: User,
    query: str,
    limit: int = 5,
    document_id: Optional[UUID] = None,
) -> list[tuple[DocumentChunk, Document]]:
    api_key = _resolve_api_key(current_user)
    if current_user.ai_provider == "google" and not api_key:
        raise HTTPException(status_code=400, detail="Google API Key not configured")

    query_vector = await get_embedding(
        query,
        is_search=True,
        provider=current_user.ai_provider,
        api_key=api_key,
    )

    statement = (
        select(DocumentChunk, Document)
        .join(Document, DocumentChunk.document_id == Document.id)
        .where(DocumentChunk.user_id == current_user.id)
    )

    if document_id:
        statement = statement.where(DocumentChunk.document_id == document_id)

    statement = statement.order_by(
        DocumentChunk.embedding.op("<=>")(query_vector)
    ).limit(limit)

    return list(session.exec(statement).all())


async def run_document_chat(
    session: Session,
    current_user: User,
    query: str,
    limit: int = 5,
    document_id: Optional[UUID] = None,
) -> tuple[str, list[tuple[DocumentChunk, Document]]]:
    matches = await search_document_chunks(
        session,
        current_user,
        query,
        limit=limit,
        document_id=document_id,
    )

    context_parts = []
    for chunk, doc in matches:
        context_parts.append(f"Documento: {doc.title}\n{chunk.content}")

    if not context_parts:
        raise HTTPException(
            status_code=404, detail="No matching document content found"
        )

    context = "\n\n".join(context_parts)

    api_key = _resolve_api_key(current_user)
    if not api_key:
        raise HTTPException(status_code=400, detail="Google API Key not configured")

    agent = build_document_chat_agent(api_key)
    deps = AgentDeps(session=session, current_user=current_user)
    prompt = (
        "Use o contexto abaixo para responder a pergunta. "
        "Se a resposta não estiver no contexto, diga que não encontrou a informação.\n\n"
        f"Contexto:\n{context}\n\nPergunta: {query}"
    )

    async with agent:
        result = await agent.run(prompt, deps=deps)

    return result.output, matches
